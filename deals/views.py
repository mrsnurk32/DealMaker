from django.shortcuts import render, redirect
from . import forms
import sqlite3
from docx import Document
from docx.shared import Inches
from .pricetemplate import TemplateMaker
def deal(request):
    if request.method == 'POST':
        form = forms.CreateDeal(request.POST)
        if form.is_valid():
            form.save()
            conn = sqlite3.connect('db.sqlite3')
            c = conn.cursor()
            for item in c.execute('Select * FROM deals_deal ORDER BY id DESC LIMIT 1'):

                id = str(item[0])
                name = item[1]
                last_name = item[2]
                price = str(item[3])
                date = str(item[4])
                time_str = str(item[5])
                time_end = str(item[6])
            deal = TemplateMaker(price)
            deal.money_str()
            x = (''.join(str(x) for x in TemplateMaker.result))

            document = Document()
            document.add_heading('Document id: %s'%(id))
            document.add_paragraph('Name: %s'%(name))
            document.add_paragraph('Lastname: %s'%(last_name))
            document.add_paragraph('price: %s (%s)00 копеек'%(price, x))
            document.add_paragraph('date: %s'%(date))
            document.add_paragraph('Start: %s'%(time_str))
            document.add_paragraph('End: %s'%(time_end))
            document.save('first document.doc')
            conn.close()

            return redirect('http://127.0.0.1:8000/')
    else:
        form = forms.CreateDeal()
    return render(request, 'deals/deal.html', {'form':form})
